"""
docx_parser.py — 解析小鹅通统一模板 PRD 文档
提取：修订记录、需求背景、章节结构、需求说明表格

Uses the underlying XML body to preserve the interleaved order of
paragraphs and tables, so each table is attached to the correct section.
"""
from __future__ import annotations

import re
from docx import Document
from docx.oxml.ns import qn


def parse_docx(filepath: str) -> dict:
    doc = Document(filepath)

    title = ""
    version = None
    author = None
    domain = None
    publish_date = None

    sections: list[dict] = []
    current_section = None
    raw_lines: list[str] = []

    # Build lookup maps for paragraphs and tables by their XML element
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    # Walk the body children in document order
    body = doc.element.body
    for child in body:
        tag = child.tag

        if tag == qn("w:p"):
            para = para_map.get(child)
            if para is None:
                continue
            text = para.text.strip()
            if not text:
                continue
            raw_lines.append(text)

            style = (para.style.name if para.style else "") or ""

            if style == "Title":
                title = text
                domain = _guess_domain(text)
                continue

            if "Heading" in style:
                if current_section:
                    sections.append(current_section)
                current_section = {"heading": text, "level": _heading_level(style), "content": [], "tables": []}
                continue

            if current_section:
                current_section["content"].append(text)

        elif tag == qn("w:tbl"):
            table = table_map.get(child)
            if table is None:
                continue
            table_data = _parse_table(table)
            if table_data:
                if current_section:
                    current_section["tables"].append(table_data)
                else:
                    sections.append({"heading": "表格", "level": 2, "content": [], "tables": [table_data]})

    if current_section:
        sections.append(current_section)

    # Extract metadata from section content
    for sec in sections:
        for line in sec["content"]:
            if not publish_date and re.search(r"发布时间[：:]", line):
                m = re.search(r"(\d{4}[/\-]\d{1,2}[/\-]\d{1,2})", line)
                if m:
                    publish_date = m.group(1).replace("/", "-")
            if not author and re.search(r"PM[：:]", line):
                m = re.search(r"PM[：:]\s*(.+)", line)
                if m:
                    author = m.group(1).strip()

    return {
        "title": title,
        "version": version,
        "author": author,
        "domain": domain,
        "publish_date": publish_date,
        "sections": sections,
        "raw_text": "\n".join(raw_lines),
    }


def _heading_level(style_name: str) -> int:
    m = re.search(r"(\d+)", style_name)
    return int(m.group(1)) if m else 2


def _guess_domain(title: str) -> str:
    title_lower = title.lower()
    if any(k in title_lower for k in ["分账", "结账", "提现"]):
        return "结账分账"
    if any(k in title_lower for k in ["红包", "转账", "零钱"]):
        return "营销资金"
    if any(k in title_lower for k in ["支付", "易宝"]):
        return "支付接入"
    if "积分" in title_lower:
        return "营销资金"
    return "其他"


def _parse_table(table) -> dict | None:
    """Parse a docx table into structured data."""
    rows = []
    headers = []
    for ri, row in enumerate(table.rows):
        # Deduplicate merged cells: docx returns the same cell object multiple times
        seen = set()
        cells = []
        for cell in row.cells:
            cell_id = id(cell._element)
            if cell_id not in seen:
                seen.add(cell_id)
                cells.append(cell.text.strip())
        if ri == 0:
            headers = cells
            continue
        rows.append(cells)

    if not rows:
        return None

    # For 需求说明 tables (原型 | 说明), extract the 说明 column
    entries = []
    desc_col = None
    for i, h in enumerate(headers):
        if h in ("说明", "描述", "需求说明", "逻辑说明"):
            desc_col = i
            break

    if desc_col is not None:
        for row_cells in rows:
            if desc_col < len(row_cells) and row_cells[desc_col].strip():
                text = row_cells[desc_col]
                items = _split_items(text)
                entries.extend(items)

    return {
        "headers": headers,
        "row_count": len(rows),
        "entries": entries if entries else [" | ".join(r) for r in rows if any(c.strip() for c in r)],
    }


def _split_items(text: str) -> list[str]:
    """Split long table cell text into logical items."""
    parts = re.split(r"\n+|(?=\d+[、.．])|(?=[①②③④⑤⑥⑦⑧⑨⑩])", text)
    result = []
    for p in parts:
        p = p.strip()
        if p and len(p) > 5:
            result.append(p)
    return result if result else [text.strip()]
